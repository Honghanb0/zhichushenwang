#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
网络安全周报生成 Agent (Cybersecurity Weekly Report Agent)  —— 核心模块
====================================================================

一个可直接运行的 Python Agent：自动采集最近 N 天全球网络安全热点事件，
抓取每篇新闻的【详细正文】，调用 DeepSeek（或任意 OpenAI 兼容）API 提炼摘要与
阅读思考，输出排版美观的 Markdown 周报。每条内容末尾自动附「来源说明 + 原文链接」。

三阶段架构：
    1. 数据采集层  -> RSS 拉取最近 N 天真实新闻
    2. 内容增强层  -> 并发抓取每篇新闻原文正文（trafilatura）
    3. 智能分析层  -> DeepSeek 大模型归类、摘要、阅读思考
    4. 报告生成层  -> Markdown 周报文件（速览表 + 热点汇总 + 阅读思考）

快速开始：
    pip install requests feedparser trafilatura
    export DEEPSEEK_API_KEY="sk-xxxxxxxx"
    python cybersecurity_weekly_agent.py

未配置 API Key 时自动回退「离线采集模式」（仍含原文正文，仅无大模型润色）。
完整参数见：python cybersecurity_weekly_agent.py --help

本模块同时被 app.py（Web 前端）导入复用。
"""

from __future__ import annotations

import argparse
import concurrent.futures
import datetime as dt
import os
import re
import sys
from dataclasses import dataclass, field
from typing import Callable, Optional

try:
    import requests
except ImportError:
    sys.exit("缺少依赖 requests，请先执行: pip install requests feedparser trafilatura")

try:
    import feedparser
except ImportError:
    sys.exit("缺少依赖 feedparser，请先执行: pip install requests feedparser trafilatura")


# ----------------------------- 配置 -----------------------------

DEFAULT_BASE_URL = "https://api.deepseek.com/v1"
DEFAULT_MODEL = "deepseek-v4-flash"            # 系统默认模型
ENV_API_KEY = "DEEPSEEK_API_KEY"               # 兼容环境变量名
REQUEST_TIMEOUT = 25
ARTICLE_CHAR_LIMIT = 1800          # 喂给大模型/离线报告的单篇正文上限
HTTP_HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; CyberSecWeeklyAgent/1.0; +https://example.com/bot)"
}

# 主流安全媒体 RSS 源（覆盖全球，免费、无需额外 API Key）
# 同时包含国内可达源（FreeBuf / 安全客 等），避免境外源不可达时整份周报为空。
DEFAULT_RSS_FEEDS = [
    # 国内网络安全资讯平台
    ("FreeBuf", "https://www.freebuf.com/feed"),
    ("安全客", "https://www.anquanke.com/rss"),
    ("嘶吼安全学院", "https://www.4hou.com/feed"),
    ("先知社区", "https://xz.aliyun.com/rss"),
    ("奇安信威胁情报中心", "https://ti.qianxin.com/blog/rss"),
    ("腾讯玄武实验室", "https://xlab.tencent.com/rss"),
    ("腾讯安全威胁情报中心", "https://s.tencent.com/rss"),
    ("阿里聚安全", "https://security.alibaba.com/blog/rss"),
    ("绿盟科技威胁情报", "https://www.nsfocus.cn/rss"),
    ("启明星辰威胁情报", "https://www.venustech.com.cn/rss"),
    ("安恒信息威胁情报", "https://www.dbappsecurity.com/rss"),
    ("看雪安全论坛", "https://bbs.pediy.com/rss.xml"),
    ("合天网安实验室", "https://www.hetianlab.com/rss"),
    ("火绒安全周报", "https://www.huorong.cn/info/rss"),
    ("知道创宇404实验室", "https://blog.knownsec.com/feed/"),
    ("长亭科技", "https://blog.chaitin.cn/feed"),
    ("微步在线威胁情报", "https://threatbook.cn/blog/rss"),
    ("天融信威胁情报中心", "https://www.topsec.com.cn/rss"),
    ("斗象科技", "https://blog.riskivy.com/feed"),
    ("威努特工控安全", "https://www.venustech-ics.com/rss"),

    # 国际一线安全新闻媒体
    ("The Hacker News", "https://feeds.feedburner.com/TheHackersNews"),
    ("BleepingComputer", "https://www.bleepingcomputer.com/feed/"),
    ("Krebs on Security", "https://krebsonsecurity.com/feed/"),
    ("Dark Reading", "https://www.darkreading.com/rss.xml"),
    ("SecurityWeek", "https://www.securityweek.com/feed/"),
    ("The Register (Security)", "https://www.theregister.com/security/headlines.atom"),
    ("Threatpost", "https://threatpost.com/feed/"),
    ("SC Magazine", "https://www.scmagazine.com/home/feed/"),
    ("Infosecurity Magazine", "https://www.infosecurity-magazine.com/rss/"),
    ("Computerworld Security", "https://www.computerworld.com/security/feed/"),
    ("ZDNet Security", "https://www.zdnet.com/topic/security/rss.xml"),
    ("Forbes Cybersecurity", "https://www.forbes.com/cybersecurity/feed/"),

    # 厂商官方安全博客 & 漏洞响应中心
    ("Google Security Blog", "https://security.googleblog.com/feeds/posts/default"),
    ("Microsoft MSRC", "https://msrc.microsoft.com/blog/feed/"),
    ("Microsoft Security Blog", "https://www.microsoft.com/en-us/security/blog/feed/"),
    ("Apple Security Updates", "https://support.apple.com/en-us/HT201222.rss"),
    ("Cloudflare Blog Security", "https://blog.cloudflare.com/tag/security/rss"),
    ("AWS Security Blog", "https://aws.amazon.com/blogs/security/feed/"),
    ("Azure Security Blog", "https://azure.microsoft.com/en-us/blog/security/feed/"),
    ("Oracle Security Blog", "https://blogs.oracle.com/security/rss"),
    ("Cisco Security Blog", "https://blogs.cisco.com/security/feed"),
    ("VMware Security Advisories", "https://www.vmware.com/security/advisories.xml"),
    ("Sophos Security Blog", "https://news.sophos.com/en-us/rss"),
    ("Palo Alto Networks Blog", "https://www.paloaltonetworks.com/blog/rss.xml"),
    ("Fortinet Blog", "https://www.fortinet.com/blog/rss.xml"),
    ("F-Secure Labs", "https://blog.f-secure.com/feed/"),
    ("ESET Security Blog", "https://www.welivesecurity.com/feed/"),
    ("Trend Micro Research", "https://www.trendmicro.com/en_us/research/rss.html"),
    ("Symantec Security Response", "https://www.broadcom.com/support/security-center/feed"),

    # 安全研究机构、情报组织、CERT
    ("SANS ISC", "https://isc.sans.edu/rssfeed_full.xml"),
    ("CISA Advisories", "https://www.cisa.gov/cybersecurity-advisories/all.xml"),
    ("CISA News", "https://www.cisa.gov/news.xml"),
    ("US-CERT", "https://www.us-cert.gov/ncas/all.xml"),
    ("NIST Cybersecurity", "https://www.nist.gov/news-events/cybersecurity/rss.xml"),
    ("ENISA", "https://www.enisa.europa.eu/news/rss"),
    ("Australian ACSC", "https://www.cyber.gov.au/rss"),
    ("Canadian CCCS", "https://www.cybercentre.canada.ca/rss.xml"),
    ("MITRE ATT&CK Updates", "https://attack.mitre.org/feed.xml"),
    ("Team Cymru", "https://team-cymru.com/feed/"),
    ("Abuse.ch", "https://abuse.ch/blog/feed/"),
    ("Malwarebytes Labs", "https://blog.malwarebytes.com/feed/"),

    # 知名安全研究员 & 技术博客
    ("Schneier on Security", "https://www.schneier.com/blog/atom.xml"),
    ("Troy Hunt", "https://www.troyhunt.com/rss"),
    ("Have I Been Pwned", "https://www.troyhunt.com/rss"),
    ("FireEye Mandiant", "https://www.mandiant.com/blog/rss.xml"),
    ("Securelist (Kaspersky)", "https://securelist.com/feed/"),
    ("Unit 42 (Palo Alto)", "https://unit42.paloaltonetworks.com/feed/"),
    ("Google TAG", "https://security.googleblog.com/feeds/posts/default/-/TAG"),
    ("Microsoft DART", "https://www.microsoft.com/en-us/security/blog/topic/dart/feed/"),
    ("Elastic Security Labs", "https://www.elastic.co/security-labs/rss"),

    # 漏洞、威胁情报专项源
    ("Exploit Database", "https://www.exploit-db.com/rss.xml"),
    ("Packet Storm Security", "https://packetstormsecurity.com/rss.xml"),
    ("VulnDB Latest Vulnerabilities", "https://vulndb.cyberriskanalytics.com/rss"),
    ("Rapid7 Blog", "https://blog.rapid7.com/rss/"),
    ("Tenable Research", "https://www.tenable.com/blog/rss.xml"),
    ("Qualys Security Blog", "https://blog.qualys.com/rss"),
]

# 默认新闻标签集合（自动打标）
DEFAULT_TAGS = ["安全资讯", "漏洞研究", "政策合规", "漏洞预警", "AI安全", "攻防对抗"]

# 默认地区选项（多选筛选）
DEFAULT_REGIONS = ["中国", "亚洲（除中国）", "欧洲", "美洲", "澳洲", "非洲"]

# RSS源地区映射（来源名 -> 地区）
RSS_REGION_MAP = {
    # 国内网络安全资讯平台 -> 中国
    "FreeBuf": "中国",
    "安全客": "中国",
    "嘶吼安全学院": "中国",
    "先知社区": "中国",
    "奇安信威胁情报中心": "中国",
    "腾讯玄武实验室": "中国",
    "腾讯安全威胁情报中心": "中国",
    "阿里聚安全": "中国",
    "绿盟科技威胁情报": "中国",
    "启明星辰威胁情报": "中国",
    "安恒信息威胁情报": "中国",
    "看雪安全论坛": "中国",
    "合天网安实验室": "中国",
    "火绒安全周报": "中国",
    "知道创宇404实验室": "中国",
    "长亭科技": "中国",
    "微步在线威胁情报": "中国",
    "天融信威胁情报中心": "中国",
    "斗象科技": "中国",
    "威努特工控安全": "中国",
    # 国际一线安全新闻媒体 -> 美洲
    "The Hacker News": "美洲",
    "BleepingComputer": "美洲",
    "Krebs on Security": "美洲",
    "Dark Reading": "美洲",
    "SecurityWeek": "美洲",
    "The Register (Security)": "欧洲",
    "Threatpost": "美洲",
    "SC Magazine": "美洲",
    "Infosecurity Magazine": "欧洲",
    "Computerworld Security": "美洲",
    "ZDNet Security": "美洲",
    "Forbes Cybersecurity": "美洲",
    # 厂商官方安全博客 -> 多地区
    "Google Security Blog": "美洲",
    "Microsoft MSRC": "美洲",
    "Microsoft Security Blog": "美洲",
    "Apple Security Updates": "美洲",
    "Cloudflare Blog Security": "美洲",
    "AWS Security Blog": "美洲",
    "Azure Security Blog": "美洲",
    "Oracle Security Blog": "美洲",
    "Cisco Security Blog": "美洲",
    "VMware Security Advisories": "美洲",
    "Sophos Security Blog": "欧洲",
    "Palo Alto Networks Blog": "美洲",
    "Fortinet Blog": "美洲",
    "F-Secure Labs": "欧洲",
    "ESET Security Blog": "欧洲",
    "Trend Micro Research": "亚洲（除中国）",
    "Symantec Security Response": "美洲",
    # 安全研究机构、情报组织、CERT -> 多地区
    "SANS ISC": "美洲",
    "CISA Advisories": "美洲",
    "CISA News": "美洲",
    "US-CERT": "美洲",
    "NIST Cybersecurity": "美洲",
    "ENISA": "欧洲",
    "Australian ACSC": "澳洲",
    "Canadian CCCS": "美洲",
    "MITRE ATT&CK Updates": "美洲",
    "Team Cymru": "欧洲",
    "Abuse.ch": "欧洲",
    "Malwarebytes Labs": "美洲",
    # 知名安全研究员 & 技术博客 -> 美洲/欧洲
    "Schneier on Security": "美洲",
    "Troy Hunt": "澳洲",
    "Have I Been Pwned": "澳洲",
    "FireEye Mandiant": "美洲",
    "Securelist (Kaspersky)": "欧洲",
    "Unit 42 (Palo Alto)": "美洲",
    "Google TAG": "美洲",
    "Microsoft DART": "美洲",
    "Elastic Security Labs": "美洲",
    # 漏洞、威胁情报专项源 -> 美洲
    "Exploit Database": "美洲",
    "Packet Storm Security": "美洲",
    "VulnDB Latest Vulnerabilities": "美洲",
    "Rapid7 Blog": "美洲",
    "Tenable Research": "美洲",
    "Qualys Security Blog": "美洲",
}
# 各标签对应的关键词（命中任一即归入该标签；都不命中则归为「安全资讯」）
TAG_KEYWORDS = {
    "漏洞研究": ["漏洞", "cve", "0day", "0-day", "exploit", "利用", "rce", "提权",
                "沙箱", "分析", "研究", "逆向", "调试", "内核", "poc"],
    "政策合规": ["法规", "合规", "政策", "监管", "条例", "gdpr", "网信办", "法律",
                "标准", "要求", "指南", "办法", "规范", "审查", "备案"],
    "漏洞预警": ["预警", "紧急", "高危", "风险", "威胁", "通报", "修复", "补丁",
                "影响", "警示", "在野", "已遭利用", "处置"],
    "AI安全": ["ai", "人工智能", "大模型", "生成式", "chatgpt", "模型", "深度伪造",
              "deepfake", "提示注入", "越狱", "llm", "机器学习", "训练数据", "幻觉"],
    "攻防对抗": ["攻击", "勒索", "钓鱼", "apt", "恶意软件", "入侵", "渗透", "僵尸网络",
                "横向移动", "ddos", "后门", "黑客", "泄露", "数据泄露", "供应链", "病毒", "木马"],
    "安全资讯": [],  # 兜底默认标签
}

# 版权声明（软件底部 / Word 页脚统一使用）
COPYRIGHT_TEXT = "【直触深网】版权所有：洪声越Jeff 联系邮箱：HongshengyueJeff@163.com"

SYSTEM_PROMPT = """你是一名资深网络安全分析师与周报主编。你的任务是将提供的安全资讯素材，整理成一份专业、结构清晰、洞察到位的中文《网络安全周报》。

写作纪律：
1. 严格基于提供的素材（含每条新闻的「标签」与「详细正文」）进行总结，不得编造素材中未出现的具体事实（如确切的 CVE 编号、厂商名称、损失金额、攻击团伙名）。信息不足时用「待核实」标注。
2. 素材已自动标注标签（安全资讯 / 漏洞研究 / 政策合规 / 漏洞预警 / AI安全 / 攻防对抗），在「详细内容」中必须保留每条事件的标签。
3. 文风专业、克制，不使用表情符号。输出纯 Markdown 文本，不要使用代码块包裹，不要添加前言客套话。"""

# 分段提示词：将「态势总览 + 详细内容」与「阅读思考」拆分为两次独立调用，
# 避免单次输出 token 上限导致内容被截断（修复「仅含阅读思考」的 bug）。
OVERVIEW_DETAIL_PROMPT = """请基于以下采集并抓取了详细正文、且已自动打标签的全球网络安全新闻素材（统计周期：{date_range}，共 {count} 条），生成周报的「一、态势总览」与「二、详细内容」两部分。

# 新闻素材
{news_block}

# 输出（纯 Markdown，不要代码块包裹）

## 一、态势总览
（3-5 段概述本周网络安全整体态势：主要风险主题、攻击与防御趋势、值得重点关注的动向。必须基于素材，不得编造。）

## 二、详细内容
（严格按素材顺序逐条展开，不得遗漏、不得合并，必须覆盖全部 {count} 条。每条保留素材给出的【标签】，格式如下：
### [标签] 事件标题
- **概述**：结合详细正文提炼该事件要点（2-4 句）。
- **时间**：YYYY-MM-DD
- **来源**：来源名
- **影响/风险**：该事件的潜在影响或风险。
- **参考链接**：URL
每条事件正文末尾另起一行附一句来源说明：
> 来源：<来源名> ｜ 原文链接：<链接>）

只输出到「二、详细内容」结束，不要包含「阅读思考」部分。"""

REFLECTION_PROMPT = """请基于以下采集并抓取了详细正文、且已自动打标签的全球网络安全新闻素材（统计周期：{date_range}，共 {count} 条），生成周报的「三、阅读思考」部分。

# 新闻素材
{news_block}

# 输出（纯 Markdown，不要代码块包裹）
## 三、阅读思考
### 1. 本周态势小结
### 2. 趋势研判
### 3. 防御建议（企业 / 个人，给出可执行项）
### 4. 后续关注

必须包含以上全部 4 个小节且内容完整，不要使用代码块包裹，不要添加前言客套话。"""


# ----------------------------- 配置对象 -----------------------------

@dataclass
class ReportConfig:
    api_key: str = ""
    base_url: str = DEFAULT_BASE_URL
    model: str = DEFAULT_MODEL
    days: int = 7
    start_date: Optional[str] = None   # 自定义开始日期 'YYYY-MM-DD'
    end_date: Optional[str] = None     # 自定义结束日期 'YYYY-MM-DD'
    max_events: int = 30
    max_full_text: int = 20          # 仅对排名前 N 的新闻抓取全文（控制耗时）
    rss_feeds: list = field(default_factory=lambda: list(DEFAULT_RSS_FEEDS))
    output: Optional[str] = None
    selected_tags: list = field(default_factory=list)      # 选中的标签筛选（空=全部）
    selected_regions: list = field(default_factory=list)    # 选中的地区筛选（空=全部）


# ----------------------------- 数据采集层 -----------------------------

def _strip_html(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def assign_region(item: dict) -> str:
    """根据新闻来源分配地区。"""
    source = item.get("source", "")
    return RSS_REGION_MAP.get(source, "美洲")  # 默认归属美洲


def _fetch_single_feed(args: tuple) -> list[dict]:
    """单线程采集单个RSS源（供并发调用）。

    设置 socket 全局超时为 REQUEST_TIMEOUT（默认20秒），
    避免个别源不可达时整个线程被永久挂起。
    """
    import socket
    source, url, start_dt, end_dt = args
    old_timeout = socket.getdefaulttimeout()
    socket.setdefaulttimeout(REQUEST_TIMEOUT)
    try:
        try:
            parsed = feedparser.parse(url)
        except Exception as exc:  # noqa: BLE001
            return [{"_error": source, "_msg": f"解析失败: {exc}"}]
        # 解析异常（HTTP 错误、解析失败等）
        if getattr(parsed, "bozo", False) and not getattr(parsed, "entries", []):
            err = getattr(parsed, "bozo_exception", "未知解析错误")
            return [{"_error": source, "_msg": f"feed 解析异常: {err}"}]
        results = []
        for entry in getattr(parsed, "entries", []):
            pub = entry.get("published_parsed") or entry.get("updated_parsed")
            if not pub:
                continue
            pub_dt = dt.datetime(*pub[:6], tzinfo=dt.timezone.utc)
            if not (start_dt <= pub_dt < end_dt):
                continue
            raw_title = entry.get("title") or ""
            title = re.sub(r"\s*https?://\S+.*$", "", raw_title)
            title = re.sub(r"\s*,\s*\([A-Za-z]{3}.*\)$", "", title).strip()
            raw_summary = _strip_html(entry.get("summary", ""))
            summary = re.sub(r"^\(c\).*?License\.?\s*", "", raw_summary).strip()
            results.append({
                "source": source,
                "title": title,
                "summary": summary[:400],
                "content": "",
                "link": entry.get("link", ""),
                "published": pub_dt,
            })
        return results
    except Exception as exc:  # noqa: BLE001
        return [{"_error": source, "_msg": str(exc)}]
    finally:
        socket.setdefaulttimeout(old_timeout)


def fetch_feed_news(start_dt: dt.datetime, end_dt: dt.datetime,
                    max_events: int, rss_feeds=None,
                    selected_regions: list = None,
                    selected_tags: list = None,
                    on_progress: Optional[Callable[[str], None]] = None) -> list[dict]:
    """并发采集 RSS 源，支持地区筛选和详细进度回调。

    selected_regions: 地区筛选列表（空列表表示全部地区）
    selected_tags: 标签筛选列表（空列表表示全部标签）
    on_progress: 进度回调函数，接收日志消息
    """
    import socket
    feeds = rss_feeds or DEFAULT_RSS_FEEDS
    
    def prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)
        print(msg, file=sys.stderr)
    
    # 地区筛选：只保留目标地区的RSS源
    target_sources = None
    if selected_regions:
        target_sources = set()
        for src, region in RSS_REGION_MAP.items():
            if region in selected_regions:
                target_sources.add(src)
        prog(f"[采集] 地区筛选：{selected_regions}，有效源数量：{len(target_sources)}")
    
    # 构建采集任务（排除不在目标地区的源）
    tasks = []
    for source, url in feeds:
        if target_sources and source not in target_sources:
            continue
        tasks.append((source, url, start_dt, end_dt))
    
    prog(f"[采集] 即将并发采集 {len(tasks)} 个RSS源，时间范围：{start_dt.date()} ~ {end_dt.date() - dt.timedelta(days=1)} ...")
    
    # 单源最大耗时（含网络+解析）。即使 socket 超时不生效，此处也强制兜底
    PER_FEED_TIMEOUT = REQUEST_TIMEOUT + 10  # 秒
    
    # 并发采集（使用线程池）
    collected: list[dict] = []
    error_sources: list = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=12) as executor:
        futures = {executor.submit(_fetch_single_feed, task): task[0] for task in tasks}
        done = 0
        total = len(futures)
        try:
            for future in concurrent.futures.as_completed(futures, timeout=PER_FEED_TIMEOUT * 2):
                source = futures[future]
                done += 1
                try:
                    results = future.result(timeout=PER_FEED_TIMEOUT)
                    for r in results:
                        if "_error" in r:
                            error_sources.append(r["_error"])
                            prog(f"  [{done}/{total}] 源 {r['_error']} 失败: {r['_msg']}")
                        else:
                            collected.append(r)
                    if done % 10 == 0 or done == total:
                        prog(f"  [采集进度] {done}/{total} 个源已完成，已采集 {len(collected)} 条")
                except concurrent.futures.TimeoutError:
                    error_sources.append(source)
                    prog(f"  [{done}/{total}] 源 {source} 超时（>{PER_FEED_TIMEOUT}秒），跳过")
                except Exception as exc:  # noqa: BLE001
                    error_sources.append(source)
                    prog(f"  [{done}/{total}] 源 {source} 异常: {exc}")
        except TimeoutError:
            # as_completed 全局超时：部分源卡死，剩余的由下面兜底清理
            prog(f"  [采集] 全局超时（>{PER_FEED_TIMEOUT * 2}秒），清理剩余 {len([f for f in futures if not f.done()])} 个未完成源")

        # 兜底：检查未完成的 future（as_completed 超时返回但未取消的）
        remaining = [f for f in futures if not f.done()]
        if remaining:
            prog(f"  [采集] {len(remaining)} 个源未在时限内完成，取消并跳过")
            for f in remaining:
                f.cancel()
                src = futures[f]
                error_sources.append(src)
                prog(f"  [采集] 取消源: {src}")
    
    prog(f"[采集] 并发采集完成，共采集 {len(collected)} 条（去除重复前），失败 {len(error_sources)} 个源）")
    
    # 为每条新闻添加地区属性
    for item in collected:
        item["region"] = assign_region(item)
    
    # 去重（标题 + 链接）
    seen: set[tuple[str, str]] = set()
    unique: list[dict] = []
    dup_count = 0
    for item in collected:
        key = (item["title"], item["link"])
        if key in seen:
            dup_count += 1
            continue
        seen.add(key)
        unique.append(item)
    
    prog(f"[去重] 去除 {dup_count} 条重复新闻，剩余 {len(unique)} 条")
    
    # 地区筛选（已在源级别筛选过，这里再做一次精确筛选）
    if selected_regions:
        before = len(unique)
        unique = [n for n in unique if n.get("region") in selected_regions]
        prog(f"[筛选] 地区筛选：{before} -> {len(unique)} 条（保留 {selected_regions}）")
    
    unique.sort(key=lambda x: x["published"], reverse=True)
    result = unique[:max_events]
    prog(f"[完成] 最终返回 {len(result)} 条新闻")
    return result


def validate_date_range(start_date: Optional[str], end_date: Optional[str],
                        days: int = 7):
    """校验自定义日期区间，返回 (start_dt, end_dt, date_range_str, error)。

    - 若同时提供 start_date 与 end_date：解析为 UTC 日界，要求 end>=start，
      结束日不超过今天，区间不超过 365 天。
    - 否则回退为最近 `days` 天。
    start_dt 为起始日 00:00 UTC；end_dt 为结束日次日 00:00 UTC（不含）。
    """
    today = dt.datetime.now(dt.timezone.utc)
    today_date = today.date()
    if start_date and end_date:
        try:
            sd = dt.datetime.strptime(start_date, "%Y-%m-%d").date()
            ed = dt.datetime.strptime(end_date, "%Y-%m-%d").date()
        except ValueError:
            return None, None, "", "日期格式应为 YYYY-MM-DD"
        if sd > ed:
            return None, None, "", "开始日期不能晚于结束日期"
        if ed > today_date:
            return None, None, "", "结束日期不能晚于今天"
        if sd > today_date:
            return None, None, "", "开始日期不能晚于今天"
        span = (ed - sd).days
        if span > 365:
            return None, None, "", "日期区间不能超过 365 天"
        start_dt = dt.datetime(sd.year, sd.month, sd.day, tzinfo=dt.timezone.utc)
        end_dt = dt.datetime(ed.year, ed.month, ed.day, tzinfo=dt.timezone.utc) \
            + dt.timedelta(days=1)
        date_range = f"{sd.strftime('%Y年%m月%d日')} - {ed.strftime('%Y年%m月%d日')}"
        return start_dt, end_dt, date_range, ""
    # 回退：最近 days 天
    end_dt = dt.datetime(today_date.year, today_date.month, today_date.day,
                         tzinfo=dt.timezone.utc) + dt.timedelta(days=1)
    start_dt = end_dt - dt.timedelta(days=days)
    sd = start_dt.date()
    ed = today_date
    date_range = f"{sd.strftime('%Y年%m月%d日')} - {ed.strftime('%Y年%m月%d日')}"
    return start_dt, end_dt, date_range, ""


# ----------------------------- 内容增强层（抓取全文） -----------------------------

def fetch_article_content(url: str, timeout: int = 20) -> str:
    """抓取单篇新闻原文正文，尝试多种提取方法。"""
    if not url or not url.startswith("http"):
        return ""
    
    # 方法1：尝试 trafilatura
    try:
        import trafilatura
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=timeout, allow_redirects=True)
        resp.raise_for_status()
        # 验证响应内容
        if len(resp.text) < 500:
            return ""
        text = trafilatura.extract(
            resp.text, include_comments=False, include_tables=False,
            include_images=False, include_links=False
        )
        if text and len(text) > 100:
            return text.strip()
    except Exception as e:
        pass
    
    # 方法2：使用 readability-lxml 后备方案
    try:
        from readability.readability import Document
        from lxml import html
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=timeout)
        if resp.status_code == 200 and len(resp.text) > 1000:
            doc = Document(resp.text)
            text = doc.summary()
            if text:
                # 去除HTML标签
                text = html.fromstring(text).text_content()
                if len(text) > 100:
                    return text.strip()
    except Exception:
        pass
    
    # 方法3：简单提取 body 标签内容
    try:
        from lxml import html
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=timeout)
        if resp.status_code == 200 and len(resp.text) > 500:
            tree = html.fromstring(resp.text)
            body = tree.find('.//body')
            if body is not None:
                text = body.text_content()
                # 清理多余空白
                import re
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) > 200:
                    return text[:5000]  # 限制长度
    except Exception:
        pass
    
    return ""


def enrich_articles(news: list[dict], max_full_text: int,
                    on_progress: Optional[Callable[[str], None]] = None) -> list[dict]:
    """并发抓取排名前 max_full_text 篇新闻的原文正文。"""
    targets = news[:max_full_text]
    
    def prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)
        print(msg, file=sys.stderr)

    def worker(item: dict) -> dict:
        copied = dict(item)
        url = copied.get("link", "")
        if url:
            copied["content"] = fetch_article_content(url)
        return copied

    prog(f"[抓取] 开始并发抓取 {len(targets)} 篇新闻原文 ...")
    
    PER_ARTICLE_TIMEOUT = 25  # 单篇最大耗时
    
    enriched = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as ex:
        futures = {ex.submit(worker, it): idx for idx, it in enumerate(targets)}
        done = 0
        try:
            for fut in concurrent.futures.as_completed(futures, timeout=None):
                idx = futures[fut]
                try:
                    result = fut.result(timeout=PER_ARTICLE_TIMEOUT)
                except concurrent.futures.TimeoutError:
                    orig = targets[idx]
                    enriched[idx] = dict(orig)
                    prog(f"  [抓取进度] {done+1}/{len(targets)} - {orig.get('source', '?')} - 超时跳过")
                    done += 1
                    continue
                except Exception as exc:  # noqa: BLE001
                    orig = targets[idx]
                    enriched[idx] = dict(orig)
                    prog(f"  [抓取进度] {done+1}/{len(targets)} - {orig.get('source', '?')} - 异常: {exc}")
                    done += 1
                    continue
                enriched[idx] = result
                done += 1
                content_len = len(result.get("content", ""))
                status = "成功" if content_len > 0 else "失败"
                prog(f"  [抓取进度] {done}/{len(targets)} - {result.get('source', '?')} - {status} - {content_len} 字符")
        except Exception as exc:  # noqa: BLE001
            prog(f"  [抓取] 全局异常: {exc}")
        
        # 兜底：取消未完成的 future
        for fut, idx in futures.items():
            if idx not in enriched and not fut.done():
                fut.cancel()
                enriched[idx] = targets[idx]

    out: list[dict] = []
    for i, item in enumerate(news):
        out.append(enriched.get(i, item))
    
    success_count = sum(1 for v in enriched.values() if len(v.get("content", "")) > 0)
    prog(f"[抓取] 完成！成功 {success_count}/{len(targets)} 篇")
    return out


# ----------------------------- 智能分析层 -----------------------------

def call_deepseek(api_key: str, base_url: str, model: str, system_prompt: str,
                  user_prompt: str, temperature: float = 0.5,
                  max_tokens: int = 6000) -> str:
    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=180)
    if resp.status_code != 200:
        raise RuntimeError(f"API 错误 {resp.status_code}: {resp.text[:500]}")
    data = resp.json()
    return data["choices"][0]["message"]["content"].strip()


def _clean_md(text: str) -> str:
    """去除模型可能包裹的代码块标记。"""
    text = (text or "").strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("markdown"):
            text = text[8:]
        text = text.strip()
    return text


# ----------------------------- 报告生成层 -----------------------------

def assign_tags(item: dict) -> list[str]:
    """基于关键词为单条新闻自动分配默认标签集合中的标签。"""
    text = " ".join([
        item.get("title", ""), item.get("summary", ""), item.get("content", "")
    ]).lower()
    tags: list[str] = []
    for tag, kws in TAG_KEYWORDS.items():
        if tag == "安全资讯":
            continue
        if any(kw in text for kw in kws):
            tags.append(tag)
    if not tags:
        tags = ["安全资讯"]
    return tags


def all_tags_of(news: list[dict]) -> list[str]:
    """返回本次搜集涉及的全部去重标签（按默认顺序排序）。"""
    seen: set[str] = set()
    for n in news:
        seen.update(assign_tags(n))
    return [t for t in DEFAULT_TAGS if t in seen]


def build_news_block(news: list[dict]) -> str:
    lines: list[str] = []
    for i, n in enumerate(news, 1):
        tags = assign_tags(n)
        tag_str = "、".join(tags)
        region = n.get("region", "美洲")
        date_str = n["published"].strftime("%Y-%m-%d")
        lines.append(f"【{i}】[{tag_str}][{region}] {n['title']}")
        lines.append(f"- 标签：{tag_str}")
        lines.append(f"- 地区：{region}")
        lines.append(f"- 来源：{n['source']}")
        lines.append(f"- 时间：{date_str}")
        lines.append(f"- 链接：{n['link']}")
        detail = n.get("content") or n.get("summary") or ""
        if detail:
            lines.append(f"- 详细内容：{detail[:ARTICLE_CHAR_LIMIT]}")
        lines.append("")
    return "\n".join(lines)


def _source_footer(n: dict) -> str:
    return f"> 来源：{n['source']} ｜ 原文链接：{n['link']}"


def _has_section(text: str, *keywords: str) -> bool:
    return all(k in text for k in keywords)


def _detail_event_blocks(body: str) -> list[str]:
    """提取「二、详细内容」下以 ### 开头的事件块列表。"""
    m = re.search(r"##\s*二、详细内容(.+?)(?=\n##\s*三|$)", body, re.S)
    section = m.group(1) if m else body
    blocks = re.split(r"(?=^###\s)", section, flags=re.M)
    return [b for b in blocks if b.strip().startswith("###")]


def _event_heading_count(body: str) -> int:
    """已覆盖（已写出）的事件条数：按「详细内容」中的 ### 事件标题计数。"""
    return len(_detail_event_blocks(body))


def _is_last_event_complete(body: str) -> bool:
    """最后一条事件是否完整（以「来源说明」引用块结尾，未被 token 截断）。"""
    blocks = _detail_event_blocks(body)
    if not blocks:
        return True
    return bool(re.search(r">\s*来源[:：]", blocks[-1]))


def _body_without_last_incomplete(body: str) -> str:
    """若末条事件被截断，截掉该不完整块，避免续写时重复。"""
    if _is_last_event_complete(body):
        return body
    blocks = _detail_event_blocks(body)
    if not blocks:
        return body
    last = blocks[-1]
    idx = body.rfind(last)
    if idx == -1:
        return body
    return body[:idx].rstrip()


def _is_body_complete(body: str, count: int) -> bool:
    """态势总览 + 详细内容 是否完整：板块齐全、覆盖全部事件且末条未截断。"""
    if not _has_section(body, "态势总览", "详细内容"):
        return False
    if count <= 0:
        return True
    covered = _event_heading_count(body)
    if covered < count:
        return False
    return _is_last_event_complete(body)


def _is_reflection_complete(reflection: str) -> bool:
    return _has_section(reflection, "本周态势小结", "趋势研判", "防御建议", "后续关注")


def _generate_body(common: dict, cfg: ReportConfig,
                   prior_body: str = "", remaining_block: str = "",
                   on_progress: Optional[Callable[[str], None]] = None) -> str:
    """生成 / 续写「一、态势总览 + 二、详细内容」。

    - prior_body 为空：首次生成，使用完整素材。
    - prior_body 非空：续写模式，要求模型在已生成内容末尾继续补充剩余事件，
      不重复已写内容（用于内容超限被截断时自动补全，忽略 token 消耗限制）。
    """
    def prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)
        print(msg, file=sys.stderr)
    
    if not prior_body:
        prog(f"[AI] 正在调用大模型生成态势总览和详细内容 ...")
        user_prompt = OVERVIEW_DETAIL_PROMPT.format(**common)
    else:
        remaining_count = remaining_block.count("\n【")
        prog(f"[AI] 内容不足，正在续写剩余 {remaining_count} 条新闻 ...")
        user_prompt = (
            "# 已生成内容（请勿重复，直接在末尾继续补充）\n"
            f"{prior_body}\n\n"
            "# 还需补充的事件素材（按素材顺序，仅以下这些）\n"
            f"{remaining_block}\n\n"
            "# 续写要求\n"
            "请仅针对「还需补充的事件素材」逐条补充「二、详细内容」中的对应条目"
            "（每条保留【标签】、概述、时间、来源、影响/风险、参考链接与来源说明），"
            "严格按素材顺序，不要重复已写内容，不要重写「一、态势总览」。"
            "直接输出这些事件的 Markdown 条目即可，不要使用代码块包裹。"
        )
    prog(f"[AI] 发送请求到 {cfg.base_url}，模型: {cfg.model}")
    result = call_deepseek(
        cfg.api_key, cfg.base_url, cfg.model, SYSTEM_PROMPT,
        user_prompt, max_tokens=8000,
    )
    prog(f"[AI] 收到响应，长度: {len(result)} 字符")
    return result


def generate_with_llm(news: list[dict], date_range: str, cfg: ReportConfig,
                     on_progress: Optional[Callable[[str], None]] = None) -> str:
    """分段生成，确保四大板块完整（修复「详细内容被截断 / 仅部分事件」的 bug）：

    - 第 1 次调用：生成 一、态势总览 + 二、详细内容（独立 token 预算）
    - 若因 token 上限导致「详细内容」未覆盖全部事件或末条被截断，
      **自动重复调用接口续写补全**（忽略 tokens 消耗限制），直至覆盖全部事件
    - 第 2 次调用：单独生成 三、阅读思考（缺失小节同样自动续写补全）
    - 「搜集时期」由程序直接注入，保证必现
    - 对空返回/被内容过滤/限流等情况回退离线生成，确保周报绝不空白
    """
    def prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)
        print(msg, file=sys.stderr)
    
    news_block = build_news_block(news)
    common = {"date_range": date_range, "count": len(news), "news_block": news_block}
    prog(f"[AI] 待处理 {len(news)} 条新闻，生成提示词长度: {len(news_block)} 字符")

    # 一、态势总览 + 二、详细内容（超限自动续写补全，忽略 token 消耗限制）
    MAX_BODY_ITER = 12
    prog(f"[AI] 第1阶段：生成态势总览和详细内容（最多 {MAX_BODY_ITER} 次调用）...")
    body = _clean_md(_generate_body(common, cfg, on_progress=prog))
    prev_covered = -1
    for i in range(MAX_BODY_ITER):
        if _is_body_complete(body, len(news)):
            prog(f"[AI] 详细内容已完整覆盖全部 {len(news)} 条新闻")
            break
        covered = _event_heading_count(body)
        if covered <= prev_covered:
            prog(f"[AI] 续写未增加覆盖条数（{covered} <= {prev_covered}），停止续写")
            break
        prev_covered = covered
        prog(f"[AI] 当前覆盖 {covered}/{len(news)} 条，需要补充 {len(news) - covered} 条")
        if not _is_last_event_complete(body):
            covered = max(0, covered - 1)  # 末条不完整，本次需重做
        remaining = news[covered:]
        if not remaining:
            break
        context = (body if _is_last_event_complete(body)
                   else _body_without_last_incomplete(body))
        cont = _clean_md(_generate_body(
            common, cfg, prior_body=context,
            remaining_block=build_news_block(remaining),
            on_progress=prog,
        ))
        cont = re.sub(r"^\s*##\s*二、详细内容\s*", "", cont)
        if not cont.strip():
            prog(f"[AI] 续写返回为空，停止续写")
            break
        body = context.rstrip() + "\n\n" + cont.lstrip()
    else:
        prog(f"[AI] 达到最大续写次数 {MAX_BODY_ITER}")
    
    prog(f"[AI] 第1阶段完成，body长度: {len(body)} 字符")

    # 三、阅读思考（缺失小节自动续写补全）
    MAX_REFLECT_ITER = 6
    prog(f"[AI] 第2阶段：生成阅读思考（最多 {MAX_REFLECT_ITER} 次调用）...")
    reflection = _clean_md(call_deepseek(
        cfg.api_key, cfg.base_url, cfg.model, SYSTEM_PROMPT,
        REFLECTION_PROMPT.format(**common), max_tokens=4000,
    ))
    for i in range(MAX_REFLECT_ITER):
        if _is_reflection_complete(reflection):
            prog(f"[AI] 阅读思考已完整")
            break
        miss = [s for s in ("本周态势小结", "趋势研判", "防御建议", "后续关注")
                if s not in reflection]
        if not miss:
            break
        prog(f"[AI] 补充缺失小节：{miss}")
        cont = _clean_md(call_deepseek(
            cfg.api_key, cfg.base_url, cfg.model, SYSTEM_PROMPT,
            "# 已生成内容（请勿重复，直接补充缺失小节）\n"
            f"{reflection}\n\n"
            "# 续写要求\n"
            f"请仅补充以下缺失的小节：{'、'.join(miss)}"
            "（格式：### 小节名 + 内容），不要重复已写内容，"
            "直接输出这些小节的 Markdown，不要使用代码块包裹。",
            max_tokens=4000,
        ))
        cont = re.sub(r"^\s*##\s*三、阅读思考\s*", "", cont)
        if not cont.strip():
            prog(f"[AI] 续写返回为空，停止续写")
            break
        reflection = reflection.rstrip() + "\n\n" + cont.lstrip()
    else:
        prog(f"[AI] 达到最大续写次数 {MAX_REFLECT_ITER}")
    
    prog(f"[AI] 第2阶段完成，reflection长度: {len(reflection)} 字符")

    # 搜集时期：程序注入，保证必现且准确
    tags_all = all_tags_of(news)

    # 兜底：大模型未产出有效内容（空返回/被内容过滤/限流/缺关键板块）时，
    # 回退离线生成，确保周报绝不空白（已采集的新闻仍会完整呈现）。
    if (not body.strip()) or (not reflection.strip()) or \
       (not _has_section(body, "详细内容")):
        return generate_offline(news, date_range)

    header = (
        f"# 网络安全周报（{date_range}）\n\n"
        f"> **搜集时期**：{date_range}（共搜集 {len(news)} 条资讯；"
        f"涉及标签：{'、'.join(tags_all)}）\n"
    )
    return header.rstrip() + "\n\n" + body.rstrip() + "\n\n" + reflection.lstrip()


def generate_offline(news: list[dict], date_range: str) -> str:
    tags_all = all_tags_of(news)
    lines: list[str] = []
    lines.append(f"# 网络安全周报（{date_range}）")
    lines.append("")
    lines.append(f"> **搜集时期**：{date_range}（共搜集 {len(news)} 条资讯；"
                 f"涉及标签：{'、'.join(tags_all)}）")
    lines.append("")
    lines.append("> 说明：本报告由 RSS 自动采集并抓取全文生成，未经过大模型润色。"
                 "在界面/参数中配置 API Key 后重新运行，可获得态势总览与阅读思考深度分析。")
    lines.append("")
    lines.append("## 一、态势总览")
    lines.append("")
    if news:
        lines.append(f"本周（{date_range}）共自动搜集 {len(news)} 条全球网络安全资讯，"
                     f"覆盖标签：{'、'.join(tags_all)}。以下为各事件详细内容，"
                     f"建议结合原文链接进一步了解事件背景与处置建议。")
    else:
        lines.append("本周未搜集到符合条件的新闻，请调整时间范围或 RSS 源后重试。")
    lines.append("")
    lines.append("## 二、详细内容")
    lines.append("")
    for i, n in enumerate(news, 1):
        tags = assign_tags(n)
        region = n.get("region", "美洲")
        date_str = n["published"].strftime("%Y-%m-%d")
        lines.append(f"### [{'、'.join(tags)}][{region}] {n['title']}")
        lines.append(f"- **标签**：{'、'.join(tags)}")
        lines.append(f"- **地区**：{region}")
        lines.append(f"- **来源**：{n['source']}")
        lines.append(f"- **时间**：{date_str}")
        lines.append(f"- **链接**：{n['link']}")
        detail = n.get("content") or n.get("summary") or ""
        if detail:
            lines.append("")
            lines.append(detail[:ARTICLE_CHAR_LIMIT])
        lines.append("")
        lines.append(_source_footer(n))
        lines.append("")
    lines.append("## 三、阅读思考")
    lines.append("")
    lines.append("本报告为自动采集模式，未包含大模型生成的深度分析与防御建议。"
                 "请配置 API Key 后重新运行以获得完整版周报。")
    return "\n".join(lines)


# ----------------------------- 统一入口 -----------------------------

def generate_report(cfg: ReportConfig,
                    on_progress: Optional[Callable[[str], None]] = None) -> str:
    def prog(msg: str) -> None:
        if on_progress:
            on_progress(msg)

    start_dt, end_dt, date_range, err = validate_date_range(
        cfg.start_date, cfg.end_date, cfg.days
    )
    if err:
        raise ValueError(f"日期范围非法：{err}")

    prog(f"[开始] 时间范围：{date_range}")
    prog("[开始] 正在并发采集 RSS 新闻 ...")
    
    news = fetch_feed_news(
        start_dt, end_dt, cfg.max_events, cfg.rss_feeds,
        selected_regions=cfg.selected_regions if cfg.selected_regions else None,
        on_progress=prog
    )
    prog(f"[采集完成] 已采集 {len(news)} 条新闻")
    
    prog("[抓取] 正在抓取原文全文 ...")
    news = enrich_articles(news, cfg.max_full_text, on_progress=prog)
    prog(f"[抓取完成] 抓取原文完成，剩余 {len(news)} 条")

    # 标签筛选（在全文抓取后进行，以便基于完整内容筛选）
    if cfg.selected_tags:
        prog(f"[标签筛选] 筛选标签：{cfg.selected_tags}")
        filtered_news = []
        for n in news:
            news_tags = assign_tags(n)
            matched = any(t in cfg.selected_tags for t in news_tags)
            prog(f"  [标签调试] 「{n.get('title', '')[:40]}...」 -> 分配标签: {news_tags}，匹配: {matched}")
            if matched:
                filtered_news.append(n)
        prog(f"[标签筛选完成] 筛选前 {len(news)} 条，筛选后 {len(filtered_news)} 条")
        news = filtered_news

    if cfg.api_key:
        if not news:
            prog("[AI] 警告：没有符合条件的新闻，跳过周报生成")
            return f"# 网络安全周报\n\n> 经过筛选，未找到符合条件的新闻。\n> 请调整筛选条件（标签/地区）或扩大时间范围后重试。\n\n{COPYRIGHT_TEXT}\n"
        prog("[AI] 正在调用大模型生成周报 ...")
        try:
            result = generate_with_llm(news, date_range, cfg, on_progress=prog)
            prog("[完成] 周报生成成功")
            return result
        except Exception as exc:  # noqa: BLE001
            prog(f"[AI] 大模型调用失败，回退离线模式：{exc}")
            return generate_offline(news, date_range)
    else:
        prog("[离线] 未配置 API Key，使用离线采集模式")
        return generate_offline(news, date_range)


def write_report(report: str, output: Optional[str] = None) -> str:
    if output:
        out_path = output
    else:
        today = dt.datetime.now().strftime("%Y-%m-%d")
        out_path = os.path.join(os.getcwd(), f"网络安全周报_{today}.md")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(report)
    return out_path


def export_docx(report_text: str, output_path: str) -> str:
    """将 Markdown 周报转换为排版美观、结构清晰的 Word(.docx) 文档。

    支持：标题层级、Markdown 表格、引用块、无序列表、加粗与行内代码、
    分隔线；正文使用中文友好字体；文末页脚附版权声明。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement

    BRAND = RGBColor(0x1F, 0x49, 0x7D)   # 深蓝
    GREY = RGBColor(0x6B, 0x72, 0x80)
    CJK_FONT = "Microsoft YaHei"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    def set_cjk(run):
        run.font.name = CJK_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), CJK_FONT)

    def add_runs(paragraph, text):
        for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("`") and part.endswith("`"):
                r = paragraph.add_run(part[1:-1])
                r.font.name = "Consolas"
                rPr = r._element.get_or_add_rPr()
                rf = rPr.find(qn("w:rFonts"))
                if rf is None:
                    rf = OxmlElement("w:rFonts")
                    rPr.append(rf)
                rf.set(qn("w:eastAsia"), "Consolas")
            else:
                r = paragraph.add_run(part)
            set_cjk(r)

    def shade_cell(cell, color="DCE6F1"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    lines = report_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 表格块：连续以 | 开头的行
        if line.lstrip().startswith("|") and i + 1 < n and lines[i + 1].lstrip().startswith("|"):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            rows = []
            for idx, tl in enumerate(tbl_lines):
                if idx == 1 and set(tl.replace("|", "").strip()) <= set("-: "):
                    continue  # 跳过分隔行
                rows.append([c.strip() for c in tl.strip().strip("|").split("|")])
            if rows:
                table = doc.add_table(rows=0, cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows):
                    cells = table.add_row().cells
                    for ci, val in enumerate(row):
                        para = cells[ci].paragraphs[0]
                        add_runs(para, val)
                        if ri == 0:
                            for r in para.runs:
                                r.bold = True
                            shade_cell(cells[ci])
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.font.size = Pt(22)
            r.font.color.rgb = BRAND
            r.bold = True
            set_cjk(r)
        elif line.startswith("## "):
            p = doc.add_heading(level=1)
            r = p.add_run(line[3:].strip())
            r.font.size = Pt(16)
            r.font.color.rgb = BRAND
            r.bold = True
            set_cjk(r)
        elif line.startswith("### "):
            p = doc.add_heading(level=2)
            r = p.add_run(line[4:].strip())
            r.font.size = Pt(13)
            r.font.color.rgb = BRAND
            r.bold = True
            set_cjk(r)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(line[2:].strip())
            r.italic = True
            r.font.color.rgb = GREY
            set_cjk(r)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
        elif line.strip() == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BFBFBF")
            pbdr.append(bottom)
            pPr.append(pbdr)
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    # 页脚版权声明
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(COPYRIGHT_TEXT)
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY
    set_cjk(fr)

    doc.save(output_path)
    return output_path


# ----------------------------- 命令行入口 -----------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="网络安全周报生成 Agent")
    parser.add_argument("--api-key", default=os.environ.get(ENV_API_KEY),
                        help=f"API Key（也可通过环境变量 {ENV_API_KEY} 传入）")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL,
                        help="API Base URL（OpenAI 兼容，如 https://api.deepseek.com/v1）")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="模型名称")
    parser.add_argument("--days", type=int, default=7, help="未指定起止日期时，统计最近 N 天（默认 7）")
    parser.add_argument("--start-date", default=None, help="自定义开始日期 YYYY-MM-DD")
    parser.add_argument("--end-date", default=None, help="自定义结束日期 YYYY-MM-DD（不超过今天）")
    parser.add_argument("--max-events", type=int, default=30, help="最多纳入事件数（默认 30）")
    parser.add_argument("--max-full-text", type=int, default=20,
                        help="抓取全文的新闻条数上限（默认 20）")
    parser.add_argument("--feeds", default=None,
                        help="自定义 RSS 源，格式：名称1|URL1,名称2|URL2")
    parser.add_argument("--output", default=None, help="输出文件路径")
    parser.add_argument("--no-llm", action="store_true",
                        help="仅使用 RSS 采集与全文抓取，不调用大模型（离线模式）")
    args = parser.parse_args()

    rss_feeds = DEFAULT_RSS_FEEDS
    if args.feeds:
        rss_feeds = []
        for part in args.feeds.split(","):
            if "|" in part:
                name, url = part.split("|", 1)
                rss_feeds.append((name.strip(), url.strip()))

    cfg = ReportConfig(
        api_key="" if args.no_llm else (args.api_key or ""),
        base_url=args.base_url,
        model=args.model,
        days=args.days,
        start_date=args.start_date,
        end_date=args.end_date,
        max_events=args.max_events,
        max_full_text=args.max_full_text,
        rss_feeds=rss_feeds,
        output=args.output,
    )

    report = generate_report(cfg, on_progress=lambda m: print(f"[*] {m}"))
    out_path = write_report(report, cfg.output)
    print(f"[+] 周报已生成：{out_path}")
    print(f"[+] 报告字数：{len(report)}")


if __name__ == "__main__":
    main()
